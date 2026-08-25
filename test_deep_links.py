#!/usr/bin/env python3
# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 James Sparenberg
"""
Unit tests for deep_links.locate_passages — the in-document passage locator
behind the Deep Links feature (design: docs/superpowers/specs/
2026-08-21-deep-links-design.md).

Pure module: no server, no Ollama. Keyword mode is fully deterministic and
tested here; semantic mode is added in a later step. Fixtures are built in a
tempdir so the tests need no corpus. Run: python3 test_deep_links.py
"""

import importlib.util
import tempfile
import zipfile
from pathlib import Path

from deep_links import locate_passages

try:
    import docx as _docx
except ImportError:
    _docx = None

_HAS_STRIPRTF = importlib.util.find_spec("striprtf") is not None
_HAS_EBOOKLIB = importlib.util.find_spec("ebooklib") is not None

try:
    from reportlab.lib.pagesizes import letter as _rl_letter
    from reportlab.pdfgen.canvas import Canvas as _RLCanvas
    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False


def _write(tmp, name, content):
    p = Path(tmp) / name
    p.write_text(content, encoding="utf-8")
    return str(p)


def test_keyword_txt_line_location_and_match_span():
    """Keyword mode on a .txt finds the passage, labels its line, marks the term."""
    body = (
        "Intro line one.\n"
        "Second paragraph here.\n"
        "The quick brown fox jumps.\n"
        "Another line about foxes.\n"
    )
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "sample.txt", body)
        res = locate_passages(path, "fox", "keyword")

    assert "passages" in res, f"expected passages, got {res!r}"
    passages = res["passages"]
    assert passages, "expected at least one passage for 'fox'"

    top = passages[0]
    # The exact word 'fox' is on line 3; 'foxes' on line 4 is a different token.
    assert top["location"] == "line 3", f"location was {top['location']!r}"

    # The match span points at the literal term inside the excerpt.
    span = top["excerpt"][top["match_start"]:top["match_end"]]
    assert span.lower() == "fox", f"match span was {span!r}"


def test_keyword_docx_section_location():
    """Keyword mode on a .docx labels the matched paragraph as 'section N'."""
    if _docx is None:
        print("SKIP: python-docx not installed")
        return
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "sample.docx")
        doc = _docx.Document()
        doc.add_paragraph("First paragraph about apples.")
        doc.add_paragraph("Second paragraph about oranges.")
        doc.add_paragraph("Third paragraph mentions the banana crop.")
        doc.add_paragraph("Fourth paragraph about pears.")
        doc.save(path)

        res = locate_passages(path, "banana", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'banana', got {res!r}"
    assert passages[0]["location"] == "section 3", \
        f"location was {passages[0]['location']!r}"
    span = passages[0]["excerpt"][
        passages[0]["match_start"]:passages[0]["match_end"]]
    assert span.lower() == "banana", f"match span was {span!r}"


def test_keyword_html_section_location():
    """Keyword mode on a .html tag-strips to text and labels blocks 'section N'."""
    html_doc = (
        "<html><head><title>T</title><style>.x{color:red}</style></head><body>"
        "<h1>Overview</h1>"
        "<p>Alpha paragraph about apples.</p>"
        "<p>Beta paragraph mentions the tornado warning.</p>"
        "<script>var x = 1;</script>"
        "<p>Gamma paragraph about pears.</p>"
        "</body></html>"
    )
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "page.html", html_doc)
        res = locate_passages(path, "tornado", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'tornado', got {res!r}"
    # blocks: Overview(1), Alpha(2), Beta+tornado(3), Gamma(4); script/style gone.
    assert passages[0]["location"] == "section 3", \
        f"location was {passages[0]['location']!r}"
    span = passages[0]["excerpt"][
        passages[0]["match_start"]:passages[0]["match_end"]]
    assert span.lower() == "tornado", f"match span was {span!r}"


def test_keyword_xml_markup_found():
    """Keyword mode on an XML/RSS file tag-strips and finds the term."""
    rss = (
        '<?xml version="1.0"?>\n<rss version="2.0"><channel>\n'
        '<title>Feed</title>\n'
        '<item><title>First</title><description>Alpha about apples.</description></item>\n'
        '<item><title>Second</title>'
        '<description>Beta mentions the hurricane alert.</description></item>\n'
        '</channel></rss>\n'
    )
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "feed.rss", rss)
        res = locate_passages(path, "hurricane", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'hurricane', got {res!r}"
    span = passages[0]["excerpt"][
        passages[0]["match_start"]:passages[0]["match_end"]]
    assert span.lower() == "hurricane", f"match span was {span!r}"
    assert passages[0]["location"].startswith("section"), passages[0]["location"]


def test_keyword_markdown_line_location():
    """Keyword mode on a .md uses the plain-text (line-based) path."""
    body = (
        "# Title\n\n"
        "Alpha paragraph.\n\n"
        "Beta mentions the avalanche risk.\n\n"
        "Gamma line.\n"
    )
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "notes.md", body)
        res = locate_passages(path, "avalanche", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'avalanche', got {res!r}"
    assert passages[0]["location"].startswith("line"), passages[0]["location"]
    span = passages[0]["excerpt"][
        passages[0]["match_start"]:passages[0]["match_end"]]
    assert span.lower() == "avalanche", f"match span was {span!r}"


def test_keyword_epub_chapter_location():
    """Keyword mode on an .epub labels passages by chapter (spine order)."""
    if not _HAS_EBOOKLIB:
        print("SKIP: ebooklib not installed")
        return
    from ebooklib import epub
    book = epub.EpubBook()
    c1 = epub.EpubHtml(title="C1", file_name="c1.xhtml")
    c1.content = "<html><body><p>Alpha chapter about apples.</p></body></html>"
    c2 = epub.EpubHtml(title="C2", file_name="c2.xhtml")
    c2.content = "<html><body><p>Beta chapter mentions the glacier survey.</p></body></html>"
    book.add_item(c1)
    book.add_item(c2)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = [c1, c2]
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "book.epub")
        epub.write_epub(path, book)
        res = locate_passages(path, "glacier", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'glacier', got {res!r}"
    assert passages[0]["location"].startswith("chapter"), passages[0]["location"]
    span = passages[0]["excerpt"][
        passages[0]["match_start"]:passages[0]["match_end"]]
    assert span.lower() == "glacier", f"match span was {span!r}"


def test_keyword_rtf_line_location():
    """Keyword mode on a .rtf finds the term; RTF has no pages, so line-based."""
    if not _HAS_STRIPRTF:
        print("SKIP: striprtf not installed")
        return
    rtf = "\n".join([
        r"{\rtf1\ansi",
        r"First line about cats.\par",
        r"Second line about the mango harvest.\par",
        r"Third line about dogs.\par",
        "}",
    ])
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "sample.rtf")
        Path(path).write_text(rtf, encoding="utf-8")
        res = locate_passages(path, "mango", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'mango', got {res!r}"
    assert passages[0]["location"] == "line 2", \
        f"location was {passages[0]['location']!r}"


_ODT_CONTENT = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<office:document-content'
    ' xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
    ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">'
    '<office:body><office:text>'
    '<text:p>First paragraph about rivers.</text:p>'
    '<text:p>Second paragraph about the volcano eruption.</text:p>'
    '<text:p>Third paragraph about mountains.</text:p>'
    '</office:text></office:body></office:document-content>'
)


def _write_odt(tmp, name):
    """Build a minimal valid .odt (mimetype + content.xml) for a fixture."""
    path = str(Path(tmp) / name)
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        zf.writestr("content.xml", _ODT_CONTENT)
    return path


def test_keyword_odt_section_location():
    """Keyword mode on a .odt labels the matched paragraph as 'section N'."""
    with tempfile.TemporaryDirectory() as tmp:
        path = _write_odt(tmp, "sample.odt")
        res = locate_passages(path, "volcano", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'volcano', got {res!r}"
    assert passages[0]["location"] == "section 2", \
        f"location was {passages[0]['location']!r}"


def test_keyword_pdf_page_location():
    """Keyword mode on a .pdf labels the passage with its page number."""
    if not _HAS_REPORTLAB:
        print("SKIP: reportlab not installed (PDF fixture writer)")
        return
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "sample.pdf")
        c = _RLCanvas(path, pagesize=_rl_letter)
        c.drawString(72, 720, "First page about kangaroos.")
        c.showPage()
        c.drawString(72, 720, "Second page mentions the platypus.")
        c.showPage()
        c.save()

        res = locate_passages(path, "platypus", "keyword")

    passages = res.get("passages")
    assert passages, f"expected a passage for 'platypus', got {res!r}"
    assert passages[0]["location"] == "p. 2", \
        f"location was {passages[0]['location']!r}"


def test_non_prose_returns_unsupported():
    """A spreadsheet extension is reported unsupported, not searched."""
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "book.xlsx", "irrelevant")
        res = locate_passages(path, "anything", "keyword")
    assert res.get("unsupported") is True, f"expected unsupported, got {res!r}"


def test_empty_doc_returns_no_passages():
    """A prose doc with no matching text yields an empty passage list."""
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "empty.txt", "nothing relevant here at all\n")
        res = locate_passages(path, "zebra", "keyword")
    assert res.get("passages") == [], f"expected [], got {res!r}"
    assert res.get("truncated") is False


def _fake_embed(texts):
    """Deterministic stand-in for Ollama: one-hot over alpha/beta/gamma markers.

    Lets the semantic path be unit-tested without a live model — cosine picks
    the text sharing the query's marker word.
    """
    out = []
    for t in texts:
        tl = t.lower()
        out.append([
            1.0 if "alpha" in tl else 0.0,
            1.0 if "beta" in tl else 0.0,
            1.0 if "gamma" in tl else 0.0,
        ])
    return out


def test_semantic_ranks_passage_and_marks_nearest_sentence():
    """Semantic mode picks the nearest passage and highlights the nearest sentence."""
    body = (
        "Intro sentence. The alpha protocol governs reactor startup. Closing remark.\n"
        "Beta procedures cover shutdown and cooling.\n"
        "Gamma notes discuss maintenance schedules.\n"
    )
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "greek.txt", body)
        res = locate_passages(path, "alpha", "semantic", embed_fn=_fake_embed)

    passages = res.get("passages")
    assert passages, f"expected a semantic passage, got {res!r}"
    top = passages[0]
    assert top["location"] == "line 1", f"location was {top['location']!r}"
    span = top["excerpt"][top["match_start"]:top["match_end"]]
    assert span == "The alpha protocol governs reactor startup.", \
        f"highlight span was {span!r}"


def test_max_passages_caps_and_flags_truncated():
    """More matches than max_passages → list is capped and truncated is True."""
    body = "".join(f"line {i} has the widget keyword\n" for i in range(5))
    with tempfile.TemporaryDirectory() as tmp:
        path = _write(tmp, "many.txt", body)
        res = locate_passages(path, "widget", "keyword", max_passages=2)
    assert len(res["passages"]) == 2, f"expected 2, got {len(res['passages'])}"
    assert res["truncated"] is True, "expected truncated=True"


def main():
    """Run every deep_links keyword-mode check; assert-based, no framework."""
    test_keyword_txt_line_location_and_match_span()
    test_keyword_html_section_location()
    test_keyword_xml_markup_found()
    test_keyword_markdown_line_location()
    test_keyword_epub_chapter_location()
    test_keyword_docx_section_location()
    test_keyword_rtf_line_location()
    test_keyword_odt_section_location()
    test_keyword_pdf_page_location()
    test_non_prose_returns_unsupported()
    test_empty_doc_returns_no_passages()
    test_semantic_ranks_passage_and_marks_nearest_sentence()
    test_max_passages_caps_and_flags_truncated()
    print("PASS: deep_links — keyword (txt/md/html/xml/docx/rtf/odt/pdf/epub) + semantic, "
          "unsupported, empty, truncation")


if __name__ == "__main__":
    main()
