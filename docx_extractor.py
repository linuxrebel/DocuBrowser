#!/usr/bin/env python3
# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 James Sparenberg
"""
DocuBrowse .docx extractor.

Uses python-docx to pull text from paragraphs and tables, plus metadata
from core properties (title, author, subject, description/keywords).

Returns a dict with the same shape as pdf_extractor.extract_pdf() so
scan_docs._extract_file() can handle it uniformly.
"""

try:
    import docx as _docx           # python-docx package (imported as ``docx``)
except ImportError:
    _docx = None


def extract_docx(file_path: str) -> dict:
    """
    Extract text and metadata from a .docx file.

    Returns:
        dict with keys: success, title, author, subject, description,
                        text (up to 5000 chars), snippet (up to 500 chars),
                        doc_type, error
    """
    result = {
        "title":       None,
        "author":      None,
        "subject":     None,
        "description": "",
        "text":        "",
        "snippet":     "",
        "doc_type":    "docx",
        "success":     False,
        "error":       None,
    }

    if _docx is None:
        result["error"] = "python-docx not installed"
        return result

    try:
        doc = _docx.Document(file_path)

        # ── Core properties (metadata) ────────────────────────────────────
        cp = doc.core_properties
        result["title"]   = (getattr(cp, "title",   "") or "").strip() or None
        result["author"]  = (getattr(cp, "author",  "") or "").strip() or None
        result["subject"] = (getattr(cp, "subject", "") or "").strip() or None

        # python-docx uses 'comments' for the dc:description field;
        # fall back to 'keywords' if comments is empty.
        comments = (getattr(cp, "comments", "") or "").strip()
        keywords = (getattr(cp, "keywords", "") or "").strip()
        result["description"] = comments or keywords or ""

        # ── Body text ─────────────────────────────────────────────────────
        parts = []

        # Paragraphs (headings, body, list items)
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Tables — each row as pipe-delimited line
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)
        result["text"]    = full_text[:5000]
        result["snippet"] = full_text[:500]

        # Mark success even on empty body so the file gets indexed
        # (metadata alone is worth recording)
        result["success"] = True

        return result

    except (OSError, ValueError, KeyError) as exc:
        result["error"] = str(exc)
        return result
